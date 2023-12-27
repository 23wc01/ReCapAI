import openai
from docx import Document
from pydub import AudioSegment
import os
import gradio
import shutil

#In cmd, set up API key with: setx "myAPIKey"
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# Transcribe audio into text
def transcribe(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe('whisper-1', audio_file)
    return transcription['text']

# Returns a dictionary of the summaries
def recap(transcription):
    summary = summarizer(transcription)
    sentiment = sentiment_analyzer(transcription)
    key_points = key_points_extracter(transcription)
    action_items = tasks_extracter(transcription)
    return {
        "sentiment": sentiment,
        "summary": summary,
        "key_points": key_points,
        "action_items": action_items
    }
    
#Analyze tone & emotion
def sentiment_analyzer(transcription):
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages = [
            {"role": "system", "content": "You are an AI experienced in language and emotion analysis, analyze the sentiment of the following text. Consider overall tone of the discussion, emotion conveyed by language used, and the context in which words and phrases are used. Indicate if sentiment is positvive, negative, or neutral, and provide brief explanations for your analysis where possible."},
            {"role": "user", "content": transcription}
        ]
    )
    return completion['choices'][0]['message']['content']

# Summarize into 1 abstract paragraph
def summarizer(transcription):
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages = [
            {"role": "system", "content": "You are a highly skilled AI trained in language comprehension and sumarization. Read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, give a coherent and readable summary that could help a person understand the main points of the disscussion without needing to read the entire text. Avoid unnecessary details or tangential points."},
            {"role": "user", "content": transcription}
        ]
    )
    return completion['choices'][0]['message']['content']

# List key points
def key_points_extracter(transcription):
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages = [
            {"role": "system", "content": "You are a proficient AI that distills information into key points. Based on the following text, identify and list the main points that were discussed. These should be the most important ideas, finidngs, or topics that are crucial to the essence of the discussion. Your goal is to provide a list someone could scan over quickly and understand what was talked about."},
            {"role": "user", "content": transcription}
        ]
    )
    return completion['choices'][0]['message']['content']

# List action items (assignments, tasks, actions, etc.)
def tasks_extracter(transcription):
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages = [
            {"role": "system", "content": "You are an AI expert in analyzing conversations and extracting action items. Please review text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. List action items clearly and concisely."},
            {"role": "user", "content": transcription}
        ]
    )
    return completion['choices'][0]['message']['content']

# Save summaries into a .docx document
def save_as_docx(recap, filename):
    doc = Document()
    for key, value in recap.items():
        #Replace all underscores in key, replace with spaces and capitalize word
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)

def printRecap(recap):
    for key, value in recap.items():
        print("\n" + key)
        print(value)
    

#Run functions
audio_file_path = "./"
audio_file_name = "EarningsCall.wav"

#Split Audio so that it's under max limit of 25 MG
audio = AudioSegment.from_mp3(audio_file_path + audio_file_name)

# Total video = 175,000 milliseconds
first_half_milliseconds = 175000/2
part1 = audio[:first_half_milliseconds]
part2 = audio[first_half_milliseconds:]
audio_file_path_part1 = audio_file_path + "Part1" + audio_file_name
audio_file_path_part2 = audio_file_path + "Part2" + audio_file_name

part1.export(audio_file_path_part1)
part2.export(audio_file_path_part2)

transcription1 = transcribe(audio_file_path_part1)
transcription2 = transcribe(audio_file_path_part2)
transcription = transcription1 + transcription2
recap_summary = recap(transcription)

printRecap(recap_summary)
save_as_docx(recap_summary, "./reacpai.docx")

def process_file(audio_filepath):
    path = "./" + os.path.basename(audio_filepath)  
    shutil.copyfile(audio_filepath.name, path)
    return recap(path)

getAudioFile = gradio.File(file_count="single", file_types=["audio"])
UI = gradio.Interface(process_file, getAudioFile, "file")
UI.launch()

    












         

