# Module to generate Microsoft Word documents
from docx import Document

# Module to generate pdf files
from fpdf import FPDF

def print_recap(recap):
    """Print formatted recap with headers & newlines
        Parameters:
        recap -- dict, Contains summary, key points, tasks, sentiment as keys and their respective content as values
    """
    for key, value in recap.items():
        # Print key names as headers & values as content after header
        print("\n" + key)
        print(value)
def save_as_txt(recap, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headers & newlines as simple text document(.txt)
        Parameters:
        recap -- dict, Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save txt to (default=./output_recap/)
        file_name -- str, name of txt (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = file_path + file_name + ".txt"
    with open(file_address, "w") as file:
        for key, value in recap.items():
            # Write key names as headers & values as content after header
            file.write("\n" + key)
            file.write("\n" + value)
    print(f"Saving {file_name}.txt at {file_path}")
    return file_address

def save_as_docx(recap, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headings & newlines as Microsoft Word docx
        Parameters:
        recap -- dict, Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save txt to (default=./output_recap/)
        file_name -- str, name of txt (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = file_path + file_name + ".docx"
    doc = Document()
    for key, value in recap.items():
        # Write key names as Heading1s in word document
        doc.add_heading(key, level=1)
        # Write values (completions) as paragraph text in word document
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(file_address)
    print(f"Saving {file_name}.docx at {file_path}")
    return file_address

def save_as_pdf(recap, file_path="./output_recap/", file_name="recapai"):
    """Save formatted recap with headings & newlines as pdf
        Parameters:
        recap -- dict, Contains summary, key points, tasks, sentiment as keys and their respective content as values
        file_path -- str, Path to save txt to (default=./output_recap/)
        file_name -- str, name of txt (default=recapai)

        Returns:
        file_address -- str, file's full path with path, name, and extension 
    """
    file_address = file_path + file_name + ".pdf"
    pdf = FPDF()
    pdf.add_page()
    for key, value in recap.items():
        # Write key names as large cornflower blue headers in pdfs
        pdf.set_font("Helvetica", style="B", size=32)
        pdf.set_text_color(100, 149, 237) # cornflower blue
        pdf.write(16, f"{key}\n")
        # Write values (completions) as paragraph text in pdf
        pdf.set_font("Helvetica",size=12)
        pdf.set_text_color(0, 0, 0) # black
        formatted_value = u'%s' % value
        pdf.write(5, formatted_value)
        pdf.write(5, "\n\n")
    pdf.output(file_address)
    print(f"Saving {file_name}.pdf to {file_path}")
    return file_address
